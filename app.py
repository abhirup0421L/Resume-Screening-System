# app.py

import streamlit as st
import streamlit.components.v1 as components
import json
import time
import os
from io import BytesIO
from docx import Document
from dotenv import load_dotenv
from backend.ai_analyzer import analyze_resume_ai
from backend.database import (
    cache_collection,
    coupons_collection,
    settings_collection
)
load_dotenv()
import plotly.graph_objects as go
from fpdf import FPDF
from streamlit_autorefresh import st_autorefresh

from backend.parser import extract_text
from backend.skills import extract_skills
from backend.matcher import match_skills
from backend.ai_validator import validate_resume
from backend.resume_generator import generate_resume

from backend.auth import (
    send_login_otp,
    verify_otp_only,
    verify_otp_and_create_account,
    login_user,
    update_name,
    get_user
)

from backend.credits import get_credits, deduct_credit
from backend.coupons import claim_coupon, create_coupon
from backend.resume_cache import (
    make_resume_hash,
    get_cached_result,
    save_cached_result
)
from backend.gemini_queue import wait_for_gemini_slot

from backend.interview import (
    LEVEL_DETAILS,
    generate_interview_questions,
    check_answers,
    update_user_progress,
    get_user_progress
)

# =====================================================
# MAINTENANCE MODE FUNCTIONS
# =====================================================

def get_maintenance_mode():

    doc = settings_collection.find_one(
        {"key": "maintenance_mode"}
    )

    return bool(
        doc and doc.get("enabled")
    )


def set_maintenance_mode(value):

    settings_collection.update_one(
        {"key": "maintenance_mode"},

        {
            "$set": {
                "enabled": value
            }
        },

        upsert=True
    )

from datetime import datetime, timedelta, timezone


# =====================================================
# MAINTENANCE MODE FUNCTIONS
# =====================================================

def get_maintenance_mode():
    doc = settings_collection.find_one({"key": "maintenance_mode"})

    if not doc or not doc.get("enabled"):
        return False, None

    end_time = doc.get("end_time")

    if end_time and datetime.now(timezone.utc) >= end_time.replace(tzinfo=timezone.utc):
        set_maintenance_mode(False)
        return False, None

    return True, end_time


def set_maintenance_mode(value, minutes=None):

    if value:

        end_time = None

        if minutes:
            end_time = (
                datetime.now(timezone.utc)
                + timedelta(minutes=minutes)
            )

        settings_collection.update_one(
            {"key": "maintenance_mode"},
            {
                "$set": {
                    "enabled": True,
                    "end_time": end_time
                }
            },
            upsert=True
        )

    else:

        settings_collection.update_one(
            {"key": "maintenance_mode"},
            {
                "$set": {
                    "enabled": False,
                    "end_time": None
                }
            },
            upsert=True
        )


def add_maintenance_time(minutes):
    doc = settings_collection.find_one({"key": "maintenance_mode"})
    now = datetime.now(timezone.utc)

    old_end = doc.get("end_time") if doc else None

    if old_end:
        old_end = old_end.replace(tzinfo=timezone.utc)
        base_time = max(old_end, now)
    else:
        base_time = now

    new_end = base_time + timedelta(minutes=minutes)

    settings_collection.update_one(
        {"key": "maintenance_mode"},
        {"$set": {"enabled": True, "end_time": new_end}},
        upsert=True
    )

# =====================================================
# PAGE CONFIG
# =====================================================
st.set_page_config(
    page_title="AI Resume Screening",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)



# =====================================================
# ADMIN PAGE
# =====================================================
if st.query_params.get("admin") == "1":

    if "admin_logged_in" not in st.session_state:
        st.session_state.admin_logged_in = False

    st.markdown("""
    <style>
    .stApp {
        background: #000000;
    }

    div[data-testid="stVerticalBlockBorderWrapper"] {
        border: 1px solid #2a2a2a !important;
        border-radius: 14px !important;
        padding: 35px !important;
        background: rgba(10,10,10,0.95) !important;
    }

    .admin-title {
        text-align: center;
        color: #ff2b2b;
        font-size: 42px;
        font-weight: 900;
        letter-spacing: 8px;
        margin-bottom: 30px;
    }

    .auth-label {
        color: white;
        font-size: 14px;
        font-weight: 800;
        margin-bottom: 8px;
        margin-top: 8px;
    }

    .stButton > button {
        background: #00ff4c;
        color: black;
        border: none;
        border-radius: 12px;
        height: 50px;
        font-weight: 900;
    }

    div[data-testid="stTextInput"] input {
        background-color: #111111 !important;
        color: white !important;
        border: 1px solid #2a2a2a !important;
        border-radius: 12px !important;
        height: 48px !important;
    }
    </style>
    """, unsafe_allow_html=True)

    left, center, right = st.columns([1.4, 1, 1.4])

    with center:
        with st.container(border=True):

            if not st.session_state.admin_logged_in:

                st.markdown(
                    '<div class="admin-title">ADMIN</div>',
                    unsafe_allow_html=True
                )

                st.markdown(
                    '<div class="auth-label">Admin ID</div>',
                    unsafe_allow_html=True
                )

                admin_id = st.text_input(
                    "",
                    placeholder="Enter admin ID",
                    key="admin_id_input",
                    label_visibility="collapsed"
                )

                st.markdown(
                    '<div class="auth-label">Password</div>',
                    unsafe_allow_html=True
                )

                admin_password = st.text_input(
                    "",
                    placeholder="Enter admin password",
                    type="password",
                    key="admin_password_input",
                    label_visibility="collapsed"
                )

                if st.button("Login", use_container_width=True):

                    if (
                        admin_id == os.getenv("ADMIN_ID")
                        and admin_password == os.getenv("ADMIN_PASSWORD")
                    ):
                        st.session_state.admin_logged_in = True
                        st.success("Admin login successful.")
                        time.sleep(1)
                        st.rerun()
                    else:
                        st.error("Invalid admin ID or password.")

                st.markdown(
                    """
                    <div style="text-align:center; margin-top:20px;">
                        <a href="/"
                           style="color:#00ff4c; text-decoration:none; font-weight:700;">
                           Back to User Login
                        </a>
                    </div>
                    """,
                    unsafe_allow_html=True
                )

                st.stop()

            st.markdown(
                '<div class="admin-title">ADMIN PANEL</div>',
                unsafe_allow_html=True
            )

            cache_count = cache_collection.count_documents({})

            st.metric("Cached Resumes", cache_count)

            st.warning("Deleting cache will remove all saved resume analysis data.")

            if st.button("🗑️ Delete All Cached Resumes", use_container_width=True):
                result = cache_collection.delete_many({})
                st.success(f"{result.deleted_count} cached resumes deleted successfully.")

            # =====================================================
            # COUPON GENERATOR
            # =====================================================

            st.write("---")
            st.subheader("🎟️ Coupon Generator")

            coupon_credits = st.number_input(
                "Credit points for coupon",
                min_value=1,
                step=1
            )

            coupon_duration = st.number_input(
                "Coupon duration in hours",
                min_value=1,
                step=1
            )

            if st.button("Generate Coupon Code", use_container_width=True):

                code = create_coupon(
                    credits=coupon_credits,
                    duration_hours=coupon_duration
                )

                st.success("Coupon created successfully.")
                st.code(code)

            active_coupons = coupons_collection.count_documents({})

            st.metric("Active Coupons", active_coupons)

            # =====================================================
            # MAINTENANCE MODE
            # =====================================================

            st.markdown("---")
            st.subheader("🚧 Maintenance Mode")

            maintenance_minutes = st.number_input(
                "Maintenance time in minutes",
                min_value=1,
                value=10,
                step=1
            )

            if st.button("🚧 Start Maintenance", use_container_width=True):
                set_maintenance_mode(True, maintenance_minutes)
                st.success("Maintenance mode started.")
                st.rerun()

            if st.button("➕ Add More Time", use_container_width=True):
                add_maintenance_time(maintenance_minutes)
                st.success("Maintenance time increased.")
                st.rerun()

            if st.button("✅ Stop Maintenance", use_container_width=True):
                set_maintenance_mode(False)
                st.success("Maintenance mode stopped.")
                st.rerun()

            if st.button("Logout", use_container_width=True):
                st.session_state.admin_logged_in = False
                st.rerun()

            st.markdown(
                """
                <div style="text-align:center; margin-top:20px;">
                    <a href="/"
                       style="color:#00ff4c; text-decoration:none; font-weight:700;">
                       Back to User Login
                    </a>
                </div>
                """,
                unsafe_allow_html=True
            )

    st.stop()




# ================================
# FULLSCREEN LOCK SCREEN
# ================================

maintenance_on, maintenance_end = get_maintenance_mode()

if maintenance_on and st.query_params.get("admin") != "1":

    remaining_text = "00:00"

    if maintenance_end:
        try:
            if maintenance_end.tzinfo is None:
                maintenance_end = maintenance_end.replace(tzinfo=timezone.utc)

            remaining = maintenance_end - datetime.now(timezone.utc)
            total_seconds = max(0, int(remaining.total_seconds()))

            hours = total_seconds // 3600
            mins = (total_seconds % 3600) // 60
            secs = total_seconds % 60

            remaining_text = f"{hours:02d}:{mins:02d}:{secs:02d}"

        except Exception:
            remaining_text = "Timer Error"

    html_code = f"""
    <style>
    .construction-overlay {{
        position: fixed;
        inset: 0;
        background: black;
        z-index: 999999;
        display: flex;
        flex-direction: column;
        justify-content: center;
        align-items: center;
    }}

    .construction-text {{
        color: red;
        font-size: 90px;
        font-weight: 900;
        letter-spacing: 6px;
        text-align: center;
    }}

    .construction-timer {{
        color: #00ff4c;
        font-size: 46px;
        font-weight: 900;
        margin-top: 35px;
        text-align: center;
    }}

    .admin-link {{
        position: fixed;
        bottom: 20px;
        right: 30px;
        color: #00ff4c !important;
        font-size: 14px;
        font-weight: 700;
        z-index: 1000000;
        text-decoration: none;
    }}
    </style>

    <div class="construction-overlay">

        <div class="construction-text">
            🚧 UNDER CONSTRUCTION 🚧
        </div>

        <div class="construction-timer">
            ⏱️ {remaining_text}
        </div>

    </div>

    <a class="admin-link" href="?admin=1">admin</a>
    """

    st.html(html_code)

    time.sleep(1)
    st.rerun()

    st.stop()

# =====================================================
# SESSION STATE
# =====================================================
defaults = {
    "page": "Resume Analyzer",
    "logged_in": False,
    "show_welcome": True,
    "user_email": "",
    "creating_account": False,
    "temp_email": "",
    "auth_mode": "login",
    "otp_sent": False,
    "otp_start_time": None,
    "otp_verified_for_signup": False,
    "verified_otp_value": "",
    "analysis_result": None,
    "admin_logged_in": False
}

for key, value in defaults.items():
    if key not in st.session_state:
        st.session_state[key] = value


# =====================================================
# CUSTOM CSS
# =====================================================
st.markdown("""
<style>

.stApp {
    background: #000000;
}

section[data-testid="stSidebar"] {
    background: #0A0A0A;
    border-right: 1px solid rgba(0,255,76,0.20);
}

div[data-testid="stVerticalBlockBorderWrapper"] {
    border: 1px solid #2a2a2a !important;
    border-radius: 14px !important;
    padding: 22px !important;
    background: rgba(10,10,10,0.95) !important;
}

.auth-small-title {
    text-align: center;
    color: white;
    font-size: 18px;
    font-weight: 800;
    margin-bottom: 10px;
}

.auth-big-title {
    text-align: center;
    color: #00ff4c;
    font-size: 38px;
    font-weight: 900;
    letter-spacing: 6px;
    margin-bottom: 18px;
}

.auth-label {
    color: white;
    font-size: 14px;
    font-weight: 800;
    margin-bottom: 6px;
    margin-top: 6px;
}

.auth-bottom {
    color: white;
    font-size: 14px;
    font-weight: 700;
    margin-top: 8px;
    margin-bottom: 8px;
}

.auth-or {
    display: flex;
    align-items: center;
    gap: 12px;
    color: #AAAAAA;
    margin: 18px 0;
    font-weight: 700;
}

.auth-or::before,
.auth-or::after {
    content: "";
    flex: 1;
    height: 1px;
    background: #2a2a2a;
}

div[data-testid="stTextInput"] {
    margin-bottom: 12px !important;
}

div[data-testid="stTextInput"] input {
    background-color: #111111 !important;
    color: white !important;
    border: 1px solid #2a2a2a !important;
    border-radius: 12px !important;
    height: 44px !important;
    box-shadow: none !important;
    outline: none !important;
}

div[data-testid="stTextInput"] input:focus {
    border: 1px solid #444444 !important;
    box-shadow: 0 0 0 1px #444444 !important;
    outline: none !important;
}

textarea {
    background-color: #111111 !important;
    color: white !important;
    border: 1px solid #2a2a2a !important;
    border-radius: 12px !important;
    box-shadow: none !important;
}

textarea:focus {
    border: 1px solid #444444 !important;
    box-shadow: 0 0 0 1px #444444 !important;
    outline: none !important;
}

div[data-baseweb="select"] > div {
    background-color: #111111 !important;
    border: 1px solid #2a2a2a !important;
    border-radius: 12px !important;
    color: white !important;
    box-shadow: none !important;
}

div[data-baseweb="select"] > div:focus-within {
    border: 1px solid #444444 !important;
    box-shadow: 0 0 0 1px #444444 !important;
}

.stButton > button {
    background: #00ff4c;
    color: black;
    border: none;
    border-radius: 12px;
    height: 46px;
    font-weight: 900;
}

.stButton > button:hover {
    background: #00cc3a;
    color: black;
}

.stDownloadButton > button {
    background: #00ff4c;
    color: black;
    border: none;
    border-radius: 12px;
    height: 46px;
    font-weight: 900;
}

.stDownloadButton > button:hover {
    background: #00cc3a;
    color: black;
}

.main-title {
    color: white;
    font-size: 38px;
    font-weight: 900;
}

.main-subtitle {
    color: #BBBBBB;
    margin-bottom: 20px;
}

.user-badge {
    position: fixed;
    top: 85px;
    left: 18px;
    z-index: 999;
    color: #00ff4c;
    font-size: 18px;
    font-weight: 800;
}

div[data-testid="stMetric"] {
    background-color: #111111;
    padding: 16px;
    border-radius: 18px;
    border: 1px solid #2a2a2a;
}

.credit-card {
    background: #111111;
    border: 1px solid #2a2a2a;
    border-radius: 14px;
    padding: 12px;
    color: #00ff4c;
    font-weight: 900;
    margin-bottom: 10px;
}

</style>
""", unsafe_allow_html=True)

# =====================================================
# WELCOME SCREEN - GREENDOT LANDING PAGE
# =====================================================

if st.session_state.show_welcome:

    if st.query_params.get("start") == "1":
        st.session_state.show_welcome = False
        st.query_params.clear()
        st.rerun()

    st.markdown("""
    <style>
    .stApp{background:#000!important;}
    header, footer, #MainMenu,
    [data-testid="stToolbar"],
    [data-testid="stDecoration"]{
        display:none!important;
    }
    .block-container{
        padding:0!important;
        margin:0!important;
        max-width:100vw!important;
    }
    iframe{
        position:fixed!important;
        inset:0!important;
        width:100vw!important;
        height:100vh!important;
        border:none!important;
    }
    </style>
    """, unsafe_allow_html=True)

    components.html("""
<!DOCTYPE html>
<html>
<head>
<style>
*{
    margin:0;
    padding:0;
    box-sizing:border-box;
}

html{
    scroll-behavior:smooth;
}

body{
    width:100vw;
    height:100vh;
    overflow-x:hidden;
    background:#111;
    font-family:Arial, sans-serif;
}

.page{
    width:100vw;
    min-height:400vh;
    position:relative;
    background:#111;
}

.section{
    width:100vw;
    height:100vh;
    position:relative;
    overflow:hidden;
}

/* ================= HOME ================= */

.home-section{
    background:
        radial-gradient(circle at 8% 95%, rgba(0,255,80,.12), transparent 30%),
        radial-gradient(circle at 80% 20%, rgba(0,255,80,.06), transparent 30%),
        linear-gradient(120deg,#06170d,#101010 45%,#191919);
}

.bg-dot{
    position:absolute;
    border-radius:50%;
    background:rgba(0,255,90,.06);
    animation:float 18s ease-in-out infinite;
}

.dot1{width:120px;height:120px;right:240px;top:0;animation-duration:20s;}
.dot2{width:90px;height:90px;right:-20px;bottom:190px;animation-duration:24s;}
.dot3{width:80px;height:80px;left:43%;bottom:50px;animation-duration:18s;}
.dot4{width:70px;height:70px;right:27%;bottom:185px;animation-duration:22s;}

@keyframes float{
    0%{transform:translate(0,0);}
    50%{transform:translate(14px,-28px);}
    100%{transform:translate(0,0);}
}

.resume-dark{
    position:absolute;
    width:310px;
    height:430px;
    left:55px;
    top:-10px;
    transform:rotate(24deg);
    background:#151515;
    border:3px solid #4ad27c;
    box-shadow:0 25px 60px rgba(0,0,0,.55);
    padding:25px;
    cursor:pointer;
}

.resume-light{
    position:absolute;
    width:330px;
    height:460px;
    left:70px;
    top:250px;
    transform:rotate(24deg);
    background:#f5fff7;
    border:3px solid #42d979;
    box-shadow:0 25px 70px rgba(0,0,0,.55);
    padding:28px;
    color:#111;
    cursor:pointer;
}

.photo{
    width:70px;
    height:95px;
    background:#45c76f;
    margin-bottom:15px;
}

.line{
    height:6px;
    background:#43d978;
    margin:9px 0;
    border-radius:10px;
}

.line.small{width:45%;}
.line.mid{width:65%;}
.line.long{width:90%;}

.grid{
    display:grid;
    grid-template-columns:repeat(2,1fr);
    gap:14px;
    margin-top:25px;
}

.box{
    height:70px;
    border:1px solid rgba(0,255,80,.5);
    border-radius:8px;
}

.circles{
    display:flex;
    flex-wrap:wrap;
    gap:7px;
    width:120px;
    margin-top:20px;
}

.circles span{
    width:18px;
    height:18px;
    border:2px solid #14b84b;
    border-radius:50%;
}

.pen{
    position:absolute;
    width:12px;
    height:210px;
    left:38px;
    bottom:100px;
    background:#111;
    border-radius:10px;
    transform:rotate(-10deg);
    box-shadow:0 0 0 2px #333;
}

.content{
    position:absolute;
    right:10%;
    top:17%;
    width:470px;
    color:white;
}

.logo{
    font-size:70px;
    font-weight:1000;
    letter-spacing:-3px;
    line-height:1;
}

.logo .green{color:#38c86b;}
.logo .white{color:white;}

.tagline{
    margin-top:22px;
    font-size:13px;
    font-weight:800;
    color:white;
}

.desc{
    margin-top:105px;
    font-size:12px;
    line-height:1.6;
    color:rgba(255,255,255,.62);
    max-width:430px;
}

.btn{
    display:inline-flex;
    align-items:center;
    justify-content:center;
    margin-top:90px;
    width:175px;
    height:45px;
    border-radius:999px;
    background:linear-gradient(135deg,#4aff85,#2ab765);
    color:#061006;
    font-size:12px;
    font-weight:1000;
    text-decoration:none;
    box-shadow:0 10px 25px rgba(0,255,80,.18);
    transition:transform .35s ease, box-shadow .35s ease, filter .35s ease;
}

.btn:hover{
    transform:translateY(-4px) scale(1.04);
    box-shadow:
        0 0 18px rgba(74,255,133,.25),
        0 18px 35px rgba(0,255,80,.22);
    filter:brightness(1.08);
}

/* ================= HOME ANIMATION ================= */

.resume-dark,
.resume-light,
.pen,
.content .logo,
.content .tagline,
.content .desc,
.content .btn{
    opacity:0;
    transition:
        opacity 1s cubic-bezier(.22,1,.36,1),
        transform 1s cubic-bezier(.22,1,.36,1);
}

.resume-dark{
    transform:rotate(24deg) translateY(-90px) translateX(-50px);
}

.resume-light{
    transform:rotate(24deg) translateY(90px) translateX(50px);
}

.pen{
    transform:rotate(-15deg) translateY(100px) translateX(60px);
}

.content .logo{transform:translateY(70px);}
.content .tagline{transform:translateY(60px);}
.content .desc{transform:translateY(50px);}
.content .btn{transform:translateY(40px);}

.home-section.home-show .resume-dark{
    opacity:1;
    transform:rotate(24deg) translateY(0) translateX(0);
}

.home-section.home-show .resume-light{
    opacity:1;
    transform:rotate(24deg) translateY(0) translateX(0);
}

.home-section.home-show .pen{
    opacity:1;
    transform:rotate(-10deg) translateY(0) translateX(0);
}

.home-section.home-show .content .logo,
.home-section.home-show .content .tagline,
.home-section.home-show .content .desc,
.home-section.home-show .content .btn{
    opacity:1;
    transform:translateY(0);
}

.home-section.home-exit .resume-dark{
    opacity:0;
    transform:rotate(24deg) translateY(-180px) translateX(-120px) scale(.9);
}

.home-section.home-exit .resume-light{
    opacity:0;
    transform:rotate(24deg) translateY(180px) translateX(120px) scale(.9);
}

.home-section.home-exit .pen{
    opacity:0;
    transform:rotate(-35deg) translateY(160px) translateX(120px) scale(.85);
}

.home-section.home-exit .content .logo{opacity:0; transform:translateY(-120px);}
.home-section.home-exit .content .tagline{opacity:0; transform:translateY(-95px);}
.home-section.home-exit .content .desc{opacity:0; transform:translateY(-75px);}
.home-section.home-exit .content .btn{opacity:0; transform:translateY(-55px);}
                    
/* ================= ABOUT ================= */

.about-section{
    background:
        radial-gradient(circle at 5% 8%, rgba(0,255,90,.07), transparent 18%),
        linear-gradient(120deg,#111,#171717);
}

.about-content{
    position:absolute;
    left:7%;
    top:19%;
    width:48%;
    color:white;
    z-index:5;

    opacity:0;
    transform:translateY(90px);

    transition:
        opacity .85s cubic-bezier(.22,1,.36,1),
        transform .85s cubic-bezier(.22,1,.36,1);
}

.about-section.show-about .about-content{
    opacity:1;
    transform:translateY(0);
}                    

.about-content h1{
    font-size:78px;
    color:#3fc46f;
    font-weight:1000;
    margin-bottom:42px;
}

.about-content p,
.about-content li{
    font-size:12px;
    line-height:1.7;
    color:rgba(255,255,255,.75);
}

.about-content h3{
    margin-top:18px;
    margin-bottom:8px;
    font-size:13px;
    color:white;
}

.about-content ul{
    padding-left:18px;
}

.bar{
    position:absolute;
    bottom:0;
    width:105px;

    border-radius:60px 60px 0 0;

    background:
        linear-gradient(
            180deg,
            #00ff55,
            #003516
        );

    box-shadow:
        0 0 35px rgba(0,255,80,.18);

    /* ANIMATION */
    transform:translateY(120%);
    transition:
        transform .9s cubic-bezier(.22,1,.36,1);
}

.about-section.show-about .bar{
    transform:translateY(0);
}

.about-section.show-about .bar1{
    transition-delay:.15s;
}

.about-section.show-about .bar2{
    transition-delay:.30s;
}

.about-section.show-about .bar3{
    transition-delay:.45s;
}                    

.bar1{height:430px;right:270px;}
.bar2{height:500px;right:150px;}
.bar3{height:610px;right:30px;}

/* ================= FEATURES ================= */

.features-section{
    background:
        radial-gradient(circle at 85% 20%, rgba(0,255,90,.08), transparent 22%),
        linear-gradient(120deg,#07120c,#151515);
}

.features-content{
    position:absolute;
    left:7%;
    top:16%;
    color:white;
}

.features-content h1{
    font-size:72px;
    color:#3fc46f;
    font-weight:1000;
    margin-bottom:45px;
}

.features-wrapper{
    display:flex;
    justify-content:space-between;
    align-items:center;
    width:88%;
    margin-top:30px;
}

.feature-tabs{
    width:42%;
}

.feature-tab{
    font-size:22px;
    font-weight:900;
    color:#00c853; /* green for all tabs */
    margin:18px 0;
    cursor:pointer;
    transition:.35s ease;
    opacity:.75;
}

/* hover effect */
.feature-tab:hover{
    color:#39ff7a;
    transform:translateX(8px);
}

/* selected tab */
.feature-tab.active{
    color:#39ff7a;
    font-size:30px; /* bigger selected tab */
    opacity:1;

    text-shadow:
        0 0 8px rgba(0,255,100,.65),
        0 0 18px rgba(0,255,100,.45),
        0 0 35px rgba(0,255,100,.25);

    transform:translateX(12px);
}

.feature-preview{
    width:42%;
    position:relative;

    left:60px;   /* move slightly right */
    top:-35px;   /* move slightly upward */
}

.quote-mark{
    position:absolute;
    top:-150px;   /* move quote upward */
    left:-10px;   /* slightly right */

    font-size:250px;
    color:rgba(0,255,100,.07);
    font-weight:1000;
}

.preview-text{
    font-size:30px;
    line-height:1.8;
    color:rgba(255,255,255,.88);
    transition:.35s ease;
}


/* ================= ABOUT EXIT ANIMATION ================= */

.about-section{
    transition:
        opacity 1.25s cubic-bezier(.22,1,.36,1),
        transform 1.25s cubic-bezier(.22,1,.36,1),
        filter 1.25s cubic-bezier(.22,1,.36,1);
}

.about-section.fade-about{
    opacity:.08;
    transform:scale(.96) translateY(-70px);
    filter:blur(7px);
}

/* bars go upward + bottom also rounded */
.about-section.fade-about .bar{
    transform:translateY(-135%);
    border-radius:60px;
    transition:1.2s cubic-bezier(.22,1,.36,1);
}

/* ================= FEATURES REVEAL ANIMATION ================= */

.features-content h1,
.feature-tab,
.feature-preview{
    opacity:0;
    transition:
        opacity 1s cubic-bezier(.22,1,.36,1),
        transform 1s cubic-bezier(.22,1,.36,1);
}

.features-content h1{
    transform:translateX(-80px);
}

.feature-tab{
    transform:translateX(-70px);
}

.feature-preview{
    transform:translateX(90px);
}

.features-section.show-features .features-content h1,
.features-section.show-features .feature-tab,
.features-section.show-features .feature-preview{
    opacity:1;
    transform:translateX(0);
}

.features-section.show-features .feature-tab:nth-child(1){transition-delay:.10s;}
.features-section.show-features .feature-tab:nth-child(2){transition-delay:.20s;}
.features-section.show-features .feature-tab:nth-child(3){transition-delay:.30s;}
.features-section.show-features .feature-tab:nth-child(4){transition-delay:.40s;}
.features-section.show-features .feature-tab:nth-child(5){transition-delay:.50s;}

.features-section.show-features .feature-preview{
    transition-delay:.35s;
}                    


/* ================= CONTACT ================= */

.features-section{
    transition:
        opacity 1.1s cubic-bezier(.22,1,.36,1),
        transform 1.1s cubic-bezier(.22,1,.36,1),
        filter 1.1s cubic-bezier(.22,1,.36,1);
}

.features-section.fade-features{
    opacity:.08;
    transform:scale(.96) translateY(-70px);
    filter:blur(7px);
}

.contact-section{
    background:
        radial-gradient(circle at 50% 0%, rgba(0,255,90,.10), transparent 30%),
        linear-gradient(120deg,#07120c,#111 55%,#151515);
}

.contact-big-title{
    position:absolute;
    top:8%;
    left:50%;
    transform:translateX(-50%) translateY(120px) scale(.92);

    font-size:130px;
    font-weight:1000;
    letter-spacing:5px;

    background:linear-gradient(
        to bottom,
        rgba(255,255,255,.20) 0%,
        rgba(255,255,255,.08) 35%,
        rgba(255,255,255,.02) 65%,
        rgba(255,255,255,0) 100%
    );

    -webkit-background-clip:text;
    -webkit-text-fill-color:transparent;
    background-clip:text;

    opacity:0;
    filter:blur(10px);
    transition:
        opacity 1.3s cubic-bezier(.22,1,.36,1),
        transform 1.3s cubic-bezier(.22,1,.36,1),
        filter 1.3s ease;
}

.contact-content{
    position:absolute;
    left:50%;
    top:42%;
    text-align:center;
    color:white;

    opacity:0;
    transform:translate(-50%,-50%) translateY(90px);

    transition:
        opacity 1.1s cubic-bezier(.22,1,.36,1),
        transform 1.1s cubic-bezier(.22,1,.36,1);
    transition-delay:.18s;
}

.contact-content h2{
    font-size:34px;
    color:rgba(255,255,255,.65);
    font-weight:900;
    margin-bottom:12px;
}

.contact-content p{
    font-size:12px;
    color:rgba(255,255,255,.32);
    margin-bottom:35px;
}

.contact-btn{
    width:190px;
    height:42px;
    margin:auto;
    border-radius:12px;
    border:1px solid rgba(255,255,255,.14);
    background:linear-gradient(135deg,rgba(255,255,255,.12),rgba(255,255,255,.04));
    backdrop-filter:blur(18px);
    display:flex;
    align-items:center;
    justify-content:space-between;
    padding:0 14px;
    color:white;
    text-decoration:none;
    font-size:12px;
    transition:.35s ease;
}

.contact-btn:hover{
    transform:translateY(-3px);
    box-shadow:0 0 25px rgba(0,255,90,.18);
}


.contact-section:not(.show-contact) .contact-content{
    transition-delay:0s;
}

.contact-section:not(.show-contact) .contact-big-title{
    transition-delay:.18s;
}
                    
.nav{
    position:fixed;
    left:50%;
    bottom:22px;
    transform:translateX(-50%);
    width:610px;
    height:58px;
    display:flex;
    align-items:center;
    padding:5px;
    border-radius:999px;
    background:linear-gradient(135deg,rgba(255,255,255,.10),rgba(255,255,255,.03));
    backdrop-filter:blur(22px);
    -webkit-backdrop-filter:blur(22px);
    border:1px solid rgba(255,255,255,.14);
    box-shadow:
        inset 0 1px 1px rgba(255,255,255,.08),
        0 15px 35px rgba(0,0,0,.28);
    z-index:9999;
    overflow:hidden;
}

.nav-slider{
    position:absolute;
    left:5px;
    top:5px;
    width:calc((100% - 10px) / 4);
    height:48px;
    border-radius:999px;
    background:linear-gradient(135deg,#00ff58,#11c65d);
    box-shadow:
        0 0 18px rgba(0,255,88,.25),
        0 8px 18px rgba(0,255,88,.18);
    transition:transform .55s cubic-bezier(.22,1,.36,1);
    z-index:0;
}

.nav a{
    flex:1;
    height:48px;
    border-radius:999px;
    color:rgba(255,255,255,.72);
    text-decoration:none;
    display:flex;
    align-items:center;
    justify-content:center;
    gap:8px;
    font-size:13px;
    font-weight:700;
    position:relative;
    z-index:2;
    background:transparent !important;
    box-shadow:none !important;
}

.nav a.active{
    color:white;
    font-weight:900;
}

.nav a:hover:not(.active){
    color:white;
}

/* ================= LIGHT THEME ================= */

.light-theme .home-section,
.light-theme .about-section,
.light-theme .features-section,
.light-theme .contact-section{
    background:linear-gradient(120deg,#f9fff9,#e7ece8);
}

.light-theme .logo .white,
.light-theme .tagline,
.light-theme .about-content,
.light-theme .features-content,
.light-theme .contact-content{
    color:#111;
}

.light-theme .desc,
.light-theme .about-content p,
.light-theme .about-content li,
.light-theme .feature-card p,
.light-theme .contact-content p{
    color:#333;
}

.light-theme .feature-card{
    background:rgba(255,255,255,.55);
    border:1px solid rgba(0,0,0,.08);
}

 /* ================= LIGHT THEME NAV FIX ================= */

.light-theme .nav{
    background:rgba(255,255,255,.68);
    backdrop-filter:blur(22px);
    -webkit-backdrop-filter:blur(22px);

    border:1px solid rgba(0,0,0,.08);

    box-shadow:
        inset 0 1px 1px rgba(255,255,255,.95),
        0 10px 30px rgba(0,0,0,.12);
}

.light-theme .nav a{
    color:rgba(0,0,0,.72);
    font-weight:700;
}

.light-theme .nav a:hover:not(.active){
    background:rgba(0,0,0,.06);
    color:#111;
}

.light-theme .nav a.active{
    background:linear-gradient(
        135deg,
        #00ff58,
        #11c65d
    );

    color:white;

    box-shadow:
        0 0 18px rgba(0,255,88,.18),
        0 8px 18px rgba(0,255,88,.14);
}

/* ================= LIGHT THEME TEXT FIX ================= */

.light-theme .preview-text{
    color:rgba(0,0,0,.78);
}

.light-theme .quote-mark{
    color:rgba(0,180,80,.12);
}

.light-theme .contact-content h2{
    color:rgba(0,0,0,.68);
}

.light-theme .contact-content p{
    color:rgba(0,0,0,.62);
}

.light-theme .contact-btn{
    color:#111;
    background:linear-gradient(
        135deg,
        rgba(0,0,0,.08),
        rgba(0,0,0,.03)
    );
    border:1px solid rgba(0,0,0,.12);
}

.light-theme .contact-big-title{
    background:linear-gradient(
        to bottom,
        rgba(0,0,0,.22) 0%,
        rgba(0,0,0,.12) 35%,
        rgba(0,0,0,.04) 65%,
        rgba(0,0,0,0) 100%
    );

    -webkit-background-clip:text;
    -webkit-text-fill-color:transparent;
    background-clip:text;
}                    

/* ================= RESPONSIVE ================= */

@media(max-width:1000px){
    .resume-dark,.resume-light,.pen{display:none;}
    .content{
        left:50%;
        right:auto;
        transform:translateX(-50%);
        text-align:center;
        width:90%;
    }
    .desc{margin:70px auto 0;}
    .nav{width:90vw;}
    .feature-grid{grid-template-columns:1fr;width:85vw;}
    .about-content{width:85%;left:8%;}
    .bar{display:none;}
}
</style>
</head>

<body id="pageBody">

<div class="page">

    <!-- HOME -->
    <section id="home" class="section home-section">

        <div class="bg-dot dot1"></div>
        <div class="bg-dot dot2"></div>
        <div class="bg-dot dot3"></div>
        <div class="bg-dot dot4"></div>

        <div class="resume-dark" id="darkResume">
            <div class="photo"></div>
            <div class="line long"></div>
            <div class="line mid"></div>
            <div class="line small"></div>
            <div class="grid">
                <div class="box"></div>
                <div class="box"></div>
                <div class="box"></div>
                <div class="box"></div>
            </div>
        </div>

        <div class="resume-light" id="lightResume">
            <div class="photo"></div>
            <div class="line long"></div>
            <div class="line mid"></div>
            <div class="line small"></div>
            <div class="grid">
                <div class="box"></div>
                <div class="box"></div>
            </div>
            <div class="circles">
                <span></span><span></span><span></span><span></span>
                <span></span><span></span><span></span><span></span>
            </div>
        </div>

        <div class="pen"></div>

        <div class="content">
            <div class="logo">
                <span class="green">Green</span><span class="white">Dot</span>
            </div>

            <div class="tagline">
                Ai Resume Analyzer + Resume Builder + Interview Qns Practice
            </div>

            <div class="desc">
                GreenDot is an AI-powered resume screening and career assistant platform designed
                to help users build, analyze, and improve resumes intelligently. It uses advanced AI
                to validate resumes, perform ATS-based skill matching, identify strengths and missing
                skills for specific job roles, and generate professional ATS-friendly resumes.
            </div>

            <a class="btn" href="?start=1">GET STARTED</a>
        </div>

    </section>

    <!-- ABOUT -->
    <section id="about" class="section about-section">

        <div class="about-content">
            <h1>ABOUT</h1>

            <p>
                GreenDot is an advanced AI-powered career platform designed to help job seekers
                create stronger resumes, improve ATS compatibility, and prepare for their dream careers.
                Our platform combines intelligent resume analysis, AI-powered CV generation, and interview
                preparation into one seamless experience.
            </p>

            <p>
                Using modern AI technology, GreenDot can analyze resumes, validate document authenticity,
                match skills with job roles, identify missing competencies, and provide smart improvement
                suggestions to increase hiring potential.
            </p>

            <h3>What We Offer</h3>

            <ul>
                <li>AI Resume Analysis – Smart ATS-based resume screening and skill matching</li>
                <li>AI CV Builder – Generate professional ATS-friendly resumes instantly</li>
                <li>Job Role Matching – Compare your skills with industry requirements</li>
                <li>Interview Preparation – Practice role-based interview questions</li>
                <li>Career Insights – Get personalized suggestions to strengthen your profile</li>
            </ul>
        </div>

        <div class="bar bar1"></div>
        <div class="bar bar2"></div>
        <div class="bar bar3"></div>

    </section>

    <!-- FEATURES -->
    <section id="features" class="section features-section">

        <div class="features-content">
            <h1>FEATURES</h1>

            <div class="features-wrapper">

                <div class="feature-tabs">

                    <div class="feature-tab active"
                    data-text="Analyze resumes with detailed AI feedback and improvement suggestions.">
                    AI Resume Analysis
                    </div>

                    <div class="feature-tab"
                    data-text="Match resumes with job roles and get an ATS compatibility score.">
                    ATS Skill Matching
                    </div>

                    <div class="feature-tab"
                    data-text="Generate professional ATS-friendly resumes for your target role.">
                    AI Resume Builder
                    </div>

                    <div class="feature-tab"
                    data-text="Practice role-based interview questions and improve confidence.">
                    Interview Practice
                    </div>

                    <div class="feature-tab"
                    data-text="Discover missing skills, strengths, and personalized suggestions.">
                    Career Insights
                    </div>

                </div>

                <div class="feature-preview">

                    <div class="quote-mark">“</div>

                    <div class="preview-text" id="previewText">
                        Analyze resumes with detailed AI feedback and improvement suggestions.
                    </div>

                </div>

            </div>
        </div>

    </section>

    <!-- CONTACT -->
    <section id="contact" class="section contact-section">

        <div class="contact-big-title">CONTACT</div>

        <div class="contact-content">
            <h2>Get In Touch</h2>

            <p>Have Questions ? Or Report any Bug ?</p>

            <a class="contact-btn" href="https://mail.google.com/mail/?view=cm&fs=1&to=abhirupd376@gmail.com&su=ResumeAI%20Support&body=Hello%20Abhirup,"
            target="_blank">
                <span>▣ Email us</span>
                <span>➜</span>
            </a>
        </div>

    </section>

    <!-- FIXED NAV -->
    <div class="nav">
        <div class="nav-slider"></div>
                    
        <a class="nav-link active" href="#home">⌂ HOME</a>
        <a class="nav-link" href="#about">ⓘ About</a>
        <a class="nav-link" href="#features">▣ Features</a>
        <a class="nav-link" href="#contact">▤ Contact</a>
    </div>

</div>

<script>
const body = document.getElementById("pageBody");
const lightResume = document.getElementById("lightResume");
const darkResume = document.getElementById("darkResume");

const links = document.querySelectorAll(".nav-link");
const sections = document.querySelectorAll(".section");

/* THEME SWITCH */
lightResume.addEventListener("click", () => {
    body.classList.add("light-theme");
});

darkResume.addEventListener("click", () => {
    body.classList.remove("light-theme");
});

/* CLICK NAV = SCROLL TO SECTION */
links.forEach(link => {
    link.addEventListener("click", function(e) {
        e.preventDefault();

        const targetId = this.getAttribute("href");
        const targetSection = document.querySelector(targetId);

        if (targetSection) {
            targetSection.scrollIntoView({
                behavior: "smooth",
                block: "start"
            });
        }

        links.forEach(l => l.classList.remove("active"));
        this.classList.add("active");
        moveSlider();
    });
});


/* MANUAL SCROLL = ACTIVE NAV CHANGE */
window.addEventListener("scroll", () => {
    let current = "";

    sections.forEach(section => {
        const sectionTop = section.offsetTop;
        const sectionHeight = section.offsetHeight;

        if (window.scrollY >= sectionTop - sectionHeight / 3) {
            current = section.getAttribute("id");
        }
    });

    links.forEach(link => {
        link.classList.remove("active");

        if (link.getAttribute("href") === "#" + current) {
            link.classList.add("active");
        }
    });
    moveSlider();
});

/* FEATURE HOVER */

const tabs =
document.querySelectorAll(".feature-tab");

const previewText =
document.getElementById("previewText");

tabs.forEach(tab => {

    tab.addEventListener("mouseenter", () => {

        tabs.forEach(t =>
            t.classList.remove("active"));

        tab.classList.add("active");

        previewText.style.opacity = "0";

        setTimeout(() => {

            previewText.innerText =
            tab.dataset.text;

            previewText.style.opacity = "1";

        }, 180);

    });

});                    

/* HOME ↔ ABOUT SMOOTH SCROLL ANIMATION */

/* HOME ↔ ABOUT ANIMATION */

const homeSection = document.getElementById("home");
const aboutSection = document.getElementById("about");

window.addEventListener("load", () => {
    setTimeout(() => {
        homeSection.classList.add("home-show");
    }, 200);
});

window.addEventListener("scroll", () => {

    const triggerPoint =
        aboutSection.offsetTop -
        window.innerHeight * 0.45;

    if(window.scrollY >= triggerPoint){

        homeSection.classList.add("home-exit");
        homeSection.classList.remove("home-show");

        aboutSection.classList.add("show-about");

    } else {

        homeSection.classList.remove("home-exit");
        homeSection.classList.add("home-show");

        aboutSection.classList.remove("show-about");
    }

});

/* ABOUT → FEATURES EXIT ANIMATION */

const featuresSection =
document.getElementById("features");

window.addEventListener("scroll", () => {

    const featureTrigger =
        featuresSection.offsetTop -
        window.innerHeight * 0.85;

    if(window.scrollY >= featureTrigger){

        aboutSection.classList.add("fade-about");

    } else {

        aboutSection.classList.remove("fade-about");
    }

});

/* FEATURES REVEAL ON SCROLL */

window.addEventListener("scroll", () => {

    const featureRevealPoint =
        featuresSection.offsetTop -
        window.innerHeight * 0.65;

    if(window.scrollY >= featureRevealPoint){
        featuresSection.classList.add("show-features");
    } else {
        featuresSection.classList.remove("show-features");
    }

});

/* FEATURES ↔ CONTACT SCROLL REVERSE ANIMATION */

const contactSection =
document.getElementById("contact");

window.addEventListener("scroll", () => {

    const start =
        contactSection.offsetTop -
        window.innerHeight;

    const end =
        contactSection.offsetTop -
        window.innerHeight * 0.20;

    let progress =
        (window.scrollY - start) /
        (end - start);

    progress = Math.max(0, Math.min(1, progress));

    /* CONTACT COMES FROM BOTTOM */
    const title =
        document.querySelector(".contact-big-title");

    const content =
        document.querySelector(".contact-content");

    title.style.opacity = progress;

    title.style.transform =
        `translateX(-50%) translateY(${120 - progress * 120}px) scale(${0.92 + progress * 0.08})`;

    title.style.filter =
        `blur(${10 - progress * 9.7}px)`;

    content.style.opacity = progress;

    content.style.transform =
        `translate(-50%,-50%) translateY(${90 - progress * 90}px)`;

    /* FEATURES PAGE GOES OUT */
    featuresSection.style.opacity =
        1 - progress * 0.88;

    featuresSection.style.transform =
        `scale(${1 - progress * 0.04})
        translateY(${-70 * progress}px)`;

    featuresSection.style.filter =
        `blur(${7 * progress}px)`;

});

/* SMOOTH NAV SLIDER */

const nav = document.querySelector(".nav");
const slider = document.querySelector(".nav-slider");

function moveSlider(){
    const active = document.querySelector(".nav-link.active");

    if(!active || !slider) return;

    const navRect = nav.getBoundingClientRect();
    const activeRect = active.getBoundingClientRect();

    const x = activeRect.left - navRect.left - 5;

    slider.style.transform = `translateX(${x}px)`;
}

window.addEventListener("load", moveSlider);
window.addEventListener("resize", moveSlider);

</script>

</body>
</html>
    """, height=950, scrolling=True)

    st.stop()

# =====================================================
# LOGIN / SIGNUP UI
# =====================================================
if not st.session_state.logged_in:

    left, center, right = st.columns([1.4, 1, 1.4])

    with center:
        with st.container(border=True):

            # ==========================
            # LOGIN PAGE
            # ==========================
            if st.session_state.auth_mode == "login":

                st.markdown(
                    '<div class="auth-small-title">Resume Screening System</div>',
                    unsafe_allow_html=True
                )
                st.markdown(
                    '<div class="auth-big-title">LOGIN</div>',
                    unsafe_allow_html=True
                )

                st.markdown(
                    '<div class="auth-label">Enter mail</div>',
                    unsafe_allow_html=True
                )

                login_email = st.text_input(
                    "",
                    placeholder="Enter your email",
                    key="login_email",
                    label_visibility="collapsed"
                )

                if login_email:
                    login_email = login_email.strip().lower()

                st.markdown(
                    '<div class="auth-label">Password</div>',
                    unsafe_allow_html=True
                )

                login_password = st.text_input(
                    "",
                    placeholder="Enter password",
                    type="password",
                    key="login_password",
                    label_visibility="collapsed"
                )

                if st.button("Login", use_container_width=True):

                    if not login_email or not login_password:
                        st.error("Please fill all fields.")

                    else:
                        success, message = login_user(
                            login_email,
                            login_password
                        )

                        if success:
                            st.session_state.logged_in = True
                            st.session_state.user_email = login_email

                            st.success(message)
                            time.sleep(1)
                            st.rerun()

                        else:
                            st.error(message)

                st.markdown(
                    '<div class="auth-or">OR</div>',
                    unsafe_allow_html=True
                )

                st.markdown(
                    '<div class="auth-bottom">Don’t have an account?</div>',
                    unsafe_allow_html=True
                )

                if st.button("Create Account", use_container_width=True):

                    st.session_state.auth_mode = "signin"
                    st.session_state.otp_sent = False
                    st.session_state.otp_start_time = None
                    st.session_state.temp_email = ""
                    st.session_state.otp_verified_for_signup = False
                    st.session_state.verified_otp_value = ""

                    st.rerun()
                
                st.markdown('<div style="text-align:center; margin-top:25px; font-size:14px;"><a href="https://mail.google.com/mail/?view=cm&fs=1&to=noname229229229@gmail.com" target="_blank" style="color:#00ff4c; text-decoration:none; font-weight:700;">Contact</a><span style="color:#444; margin:0 10px;"> | </span><a href="?admin=1" style="color:#00ff4c; text-decoration:none; font-weight:700;">Admin</a></div>', unsafe_allow_html=True)


            # ==========================
            # SIGN IN / CREATE ACCOUNT
            # ==========================
            elif st.session_state.auth_mode == "signin":

                st.markdown(
                    '<div class="auth-small-title">Resume Screening System</div>',
                    unsafe_allow_html=True
                )
                st.markdown(
                    '<div class="auth-big-title">SIGN IN</div>',
                    unsafe_allow_html=True
                )

                st.markdown(
                    '<div class="auth-label">Enter mail</div>',
                    unsafe_allow_html=True
                )

                signin_email = st.text_input(
                    "",
                    placeholder="Enter your email",
                    key="signin_email",
                    label_visibility="collapsed",
                    disabled=st.session_state.otp_sent
                )

                if signin_email:
                    signin_email = signin_email.strip().lower()

                existing_user = None

                if signin_email:
                    existing_user = get_user(signin_email)

                account_exists = bool(
                    existing_user
                    and existing_user.get("verified")
                    and existing_user.get("password_hash")
                )

                if account_exists and not st.session_state.otp_sent:
                    st.warning("⚠️ Account already exists")

                if st.button(
                    "Send OTP",
                    use_container_width=True,
                    disabled=bool(st.session_state.otp_sent or (account_exists and not st.session_state.otp_sent))
                ):

                    if signin_email:

                        with st.spinner("Sending OTP..."):
                            sent = send_login_otp(signin_email)

                        if sent:
                            st.session_state.temp_email = signin_email
                            st.session_state.otp_sent = True
                            st.session_state.otp_start_time = time.time()

                            st.success("OTP sent successfully.")
                            st.rerun()

                        else:
                            st.error("Failed to send OTP.")

                    else:
                        st.error("Please enter email.")

                # ==========================
                # OTP SECTION
                # ==========================
                if st.session_state.otp_sent:

                    # refresh ONLY until OTP is verified
                    if not st.session_state.otp_verified_for_signup:
                        st_autorefresh(interval=1000, key="otp_refresh")

                    st.markdown(
                        '<div class="auth-label">Enter OTP</div>',
                        unsafe_allow_html=True
                    )

                    otp = st.text_input(
                        "",
                        placeholder="Enter OTP",
                        key="signup_otp",
                        label_visibility="collapsed",
                        disabled=st.session_state.otp_verified_for_signup
                    )

                    if otp:
                        otp = otp.strip()

                    if st.session_state.otp_verified_for_signup:
                        st.success("OTP verified. Now set your password.")

                    else:
                        elapsed = int(time.time() - st.session_state.otp_start_time)
                        remaining = max(0, 300 - elapsed)

                        mins = remaining // 60
                        secs = remaining % 60

                        st.markdown(
                            f"""
                            <div style="
                                background:#102033;
                                color:#00ff4c;
                                padding:12px;
                                border-radius:10px;
                                margin-top:-8px;
                                margin-bottom:18px;
                                font-weight:700;
                            ">
                                OTP expires in: {mins:02d}:{secs:02d}
                            </div>
                            """,
                            unsafe_allow_html=True
                        )

                        if remaining == 0:
                            st.session_state.otp_sent = False
                            st.session_state.otp_start_time = None
                            st.session_state.temp_email = ""
                            st.session_state.otp_verified_for_signup = False
                            st.session_state.verified_otp_value = ""

                            st.error("OTP expired. Please send OTP again.")
                            st.rerun()

                        # AUTO VERIFY OTP
                        if (
                            otp
                            and len(otp) == 6
                            and not st.session_state.otp_verified_for_signup
                        ):

                            success, message = verify_otp_only(
                                st.session_state.temp_email,
                                otp
                            )

                            if success:
                                st.session_state.otp_verified_for_signup = True
                                st.session_state.verified_otp_value = otp
                                st.session_state.otp_start_time = None

                                st.success("OTP verified. Now set your password.")
                                st.rerun()

                            else:
                                st.error(message)

                    # ==========================
                    # PASSWORD SECTION
                    # ==========================
                    if st.session_state.otp_verified_for_signup:

                        st.markdown(
                            '<div class="auth-label">Set Password</div>',
                            unsafe_allow_html=True
                        )

                        new_password = st.text_input(
                            "",
                            placeholder="Set password",
                            type="password",
                            key="new_password",
                            label_visibility="collapsed"
                        )

                        confirm_password = st.text_input(
                            "",
                            placeholder="Confirm password",
                            type="password",
                            key="confirm_password",
                            label_visibility="collapsed"
                        )

                        if st.button(
                            "Create Account",
                            use_container_width=True,
                            disabled=st.session_state.get("creating_account", False)
                        ):

                            st.session_state.creating_account = True

                            if not new_password or not confirm_password:

                                st.session_state.creating_account = False
                                st.error("Please fill password fields.")

                            elif new_password != confirm_password:

                                st.session_state.creating_account = False
                                st.error("Passwords do not match.")

                            else:

                                success, message = verify_otp_and_create_account(
                                    st.session_state.temp_email,
                                    st.session_state.verified_otp_value,
                                    new_password
                                )

                                if success:

                                    st.session_state.logged_in = True
                                    st.session_state.user_email = st.session_state.temp_email

                                    st.session_state.otp_sent = False
                                    st.session_state.otp_start_time = None
                                    st.session_state.temp_email = ""
                                    st.session_state.otp_verified_for_signup = False
                                    st.session_state.verified_otp_value = ""

                                    st.success(message)

                                    st.session_state.creating_account = False

                                    time.sleep(1)
                                    st.rerun()

                                else:

                                    st.session_state.creating_account = False
                                    st.error(message)

                st.markdown(
                    '<div class="auth-or">OR</div>',
                    unsafe_allow_html=True
                )

                st.markdown(
                    '<div class="auth-bottom">Already have an account?</div>',
                    unsafe_allow_html=True
                )

                if st.button("Back to Login", use_container_width=True):

                    st.session_state.auth_mode = "login"
                    st.session_state.otp_sent = False
                    st.session_state.otp_start_time = None
                    st.session_state.temp_email = ""
                    st.session_state.otp_verified_for_signup = False
                    st.session_state.verified_otp_value = ""

                    st.rerun()

    st.stop()


# =====================================================
# MAIN USER
# =====================================================
user = get_user(st.session_state.user_email)

display_name = "User"

if user and user.get("name"):
    display_name = user.get("name")

credits = get_credits(st.session_state.user_email)

st.markdown(
    f'<div class="user-badge">👤 {display_name}</div>',
    unsafe_allow_html=True
)


# =====================================================
# SIDEBAR
# =====================================================
with st.sidebar:

    st.title("⚡ Resume AI")
    st.caption("AI Resume Analyzer & CV Maker")

    st.write("---")

    # COUPON SYSTEM
    coupon_code = st.text_input(
        "Coupon Code",
        placeholder="Enter coupon code"
    )

    if st.button("🎁 Claim Coupon", use_container_width=True):

        if not coupon_code:
            st.error("Please enter coupon code.")

        else:

            success, message = claim_coupon(
                st.session_state.user_email,
                coupon_code
            )

            if success:
                st.success(message)
                time.sleep(1)
                st.rerun()

            else:
                st.error(message)

    st.write("---")

    current_name = ""

    if user:
        current_name = user.get("name", "")

    profile_name = st.text_input("Name", value=current_name)

    if st.button("💾 Save Name", use_container_width=True):
        update_name(st.session_state.user_email, profile_name)
        st.success("Name updated.")

    st.write("📧", st.session_state.user_email)

    if st.button("🚪 Logout", use_container_width=True):

        st.session_state.logged_in = False
        st.session_state.user_email = ""
        st.session_state.temp_email = ""
        st.session_state.auth_mode = "login"
        st.session_state.otp_sent = False
        st.session_state.otp_start_time = None
        st.session_state.otp_verified_for_signup = False
        st.session_state.verified_otp_value = ""
        st.session_state.analysis_result = None

        st.rerun()

    st.markdown(
        """
        <div style="
            position: fixed;
            bottom: 10px;
            left: 25px;
            color: #666666;
            font-size: 12px;
            font-weight: 600;
        ">
            v3.0.0 UI Updated
        </div>
        """,
        unsafe_allow_html=True
    )


# =====================================================
# HEADER
# =====================================================
header_left, header_right = st.columns([6, 1])

with header_left:
    st.markdown(
        '<div class="main-title">📄 AI Resume Screening System</div>',
        unsafe_allow_html=True
    )

with header_right:
    st.markdown(
        f"""
        <div style="
            background:#111111;
            border:1px solid #2a2a2a;
            padding:12px 20px;
            border-radius:14px;
            color:#00ff4c;
            font-weight:900;
            font-size:18px;
            text-align:center;
            margin-top:10px;
        ">
            💳 {credits} Credits
        </div>
        """,
        unsafe_allow_html=True
    )

st.markdown(
    '<div class="main-subtitle">Smart ATS Analyzer + AI CV Maker Powered by Gemini AI</div>',
    unsafe_allow_html=True
)
st.warning("⚠️ Saved resume analysis data is automatically cleared after 10 minutes.")

nav1, nav2, nav3 = st.columns(3)

with nav1:
    if st.button("🔍 Resume Analyzer", use_container_width=True):
        st.session_state.page = "Resume Analyzer"
        st.rerun()

with nav2:
    if st.button("✨ AI CV Maker", use_container_width=True):
        st.session_state.page = "AI CV Maker"
        st.rerun()

with nav3:
    if st.button(
        "🎯 Interview Prep",
        use_container_width=True
    ):
        st.session_state.page = "Interview Prep"
        st.rerun()

st.write("---")


# =====================================================
# LOAD JOB ROLES
# =====================================================
try:
    with open("data/roles.json", "r") as file:
        roles_data = json.load(file)
except FileNotFoundError:
    st.error("❌ data/roles.json file not found.")
    st.stop()

job_roles = list(roles_data.keys())

if not job_roles:
    st.error("❌ No roles found in data/roles.json.")
    st.stop()


# =====================================================
# PDF FUNCTION
# =====================================================
def create_pdf(resume_text):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=14)

    pdf.set_fill_color(20, 20, 20)
    pdf.rect(0, 0, 210, 297, "F")

    pdf.set_text_color(0, 255, 76)
    pdf.set_font("Arial", "B", 20)
    pdf.cell(0, 12, "AI GENERATED RESUME", ln=True, align="C")

    pdf.ln(6)

    pdf.set_text_color(255, 255, 255)
    pdf.set_font("Arial", size=10)

    for line in resume_text.split("\n"):
        clean_line = line.encode("latin-1", "replace").decode("latin-1")

        if clean_line.strip().isupper() and len(clean_line.strip()) < 35:
            pdf.ln(3)
            pdf.set_text_color(0, 255, 76)
            pdf.set_font("Arial", "B", 12)
            pdf.multi_cell(0, 8, clean_line)
            pdf.set_text_color(255, 255, 255)
            pdf.set_font("Arial", size=10)
        else:
            pdf.multi_cell(0, 6, clean_line)

    return pdf.output(dest="S").encode("latin-1")

def create_docx(resume_text):
    doc = Document()

    for line in resume_text.split("\n"):
        clean_line = line.strip()

        if not clean_line:
            doc.add_paragraph("")
            continue

        if clean_line.isupper() and len(clean_line) < 40:
            doc.add_heading(clean_line, level=2)
        else:
            doc.add_paragraph(clean_line)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer

def show_resume_preview(
    resume_text,
    add_photo_space="No",
    add_certificate_space="No"
):

    safe_text = (
        resume_text
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )

    lines = safe_text.split("\n")
    html_lines = []

    in_list = False

    for line in lines:

        clean = line.strip()

        if not clean:
            continue

        # Ignore AI placeholders
        if clean in [
            "[Attach Photo Here]",
            "[Attach Certificates Here]",
            "PHOTO SPACE",
            "CERTIFICATES SPACE"
        ]:
            continue

        elif clean.isupper() and len(clean) < 40:

            if in_list:
                html_lines.append("</ul>")
                in_list = False

            html_lines.append(
                f"""
                <h2 style="
                    color:#00aa33;
                    margin-top:24px;
                    border-bottom:1px solid #cccccc;
                    padding-bottom:6px;
                ">
                    {clean}
                </h2>
                """
            )

        elif clean.startswith("-"):

            if not in_list:
                html_lines.append("<ul>")
                in_list = True

            html_lines.append(
                f"<li style='margin-bottom:8px;'>{clean[1:].strip()}</li>"
            )

        else:

            if in_list:
                html_lines.append("</ul>")
                in_list = False

            html_lines.append(
                f"<p style='margin:8px 0;'>{clean}</p>"
            )

    if in_list:
        html_lines.append("</ul>")

    html_content = "\n".join(html_lines)

    # PHOTO BOX (RIGHT SIDE)
    photo_box = ""

    if add_photo_space == "Yes":
        photo_box = """
        <div style="
            width:150px;
            height:180px;
            border:2px dashed #00aa33;
            border-radius:10px;
            display:flex;
            align-items:center;
            justify-content:center;
            text-align:center;
            color:#555;
            font-weight:700;
            flex-shrink:0;
        ">
            Attach<br>Photo<br>Here
        </div>
        """

    # CERTIFICATE BOX
    certificate_box = ""

    if add_certificate_space == "Yes":
        certificate_box = """
        <div style="
            margin-top:30px;
            border:2px dashed #00aa33;
            border-radius:12px;
            height:180px;
            display:flex;
            align-items:center;
            justify-content:center;
            color:#555;
            font-size:18px;
            font-weight:700;
        ">
            Attach Certificates Here
        </div>
        """

    st.html(f"""
    <div style="
        background:white;
        color:#111111;
        max-width:850px;
        margin:auto;
        padding:45px;
        border-radius:22px;
        box-shadow:0 0 35px rgba(0,255,76,0.22);
        font-family:Arial,sans-serif;
        line-height:1.6;
    ">

    <h1 style="
        text-align:center;
        margin-bottom:25px;
        color:#111111;
    ">
    Professional Resume
    </h1>

    <div style="
        display:flex;
        justify-content:space-between;
        align-items:flex-start;
        gap:30px;
    ">

        <div style="flex:1;">
            {html_content}
        </div>

        {photo_box}

    </div>

    {certificate_box}

    </div>
    """)

# =====================================================
# RESULT DISPLAY FUNCTION
# =====================================================
def show_analysis_result(ai, score, matched, missing):

    st.write("---")

    st.subheader("👤 Candidate Details")

    a, b, c = st.columns(3)

    with a:
        st.metric("Name", ai.get("candidate_name", "Unknown"))

    with b:
        st.metric("Experience", ai.get("experience_level", "Unknown"))

    with c:
        st.metric("AI Confidence", f'{ai.get("confidence", 0)}%')

    st.info(ai.get("summary", ""))

    if ai.get("ai_summary"):
        st.subheader("🤖 AI Resume Analysis")
        st.info(ai.get("ai_summary"))

    if ai.get("role_fit"):
        st.metric("Role Fit", ai.get("role_fit"))

    if ai.get("strengths"):
        st.subheader("💪 Strengths")
        for item in ai.get("strengths", []):
            st.success(item)

    if ai.get("weaknesses"):
        st.subheader("⚠️ Weaknesses")
        for item in ai.get("weaknesses", []):
            st.warning(item)

    if ai.get("improvement_suggestions"):
        st.subheader("🚀 Improvement Suggestions")
        for item in ai.get("improvement_suggestions", []):
            st.info(item)

    st.write("---")

    st.subheader("📊 ATS Score")

    s1, s2 = st.columns(2)

    with s1:
        st.metric("Match Score", f"{score}%")
        st.progress(score / 100)

    with s2:
        if score >= 85:
            st.success("Excellent Match")
        elif score >= 65:
            st.warning("Good Match")
        else:
            st.error("Needs Improvement")

    st.subheader("📈 Skill Breakdown")

    fig = go.Figure(data=[go.Pie(
        labels=["Matched", "Missing"],
        values=[len(matched), len(missing)],
        hole=0.55
    )])

    st.plotly_chart(fig, use_container_width=True)

    x, y = st.columns(2)

    with x:
        st.subheader("✅ Matched Skills")

        if matched:
            for skill in matched:
                st.success(skill.title())
        else:
            st.info("No matched skills found.")

    with y:
        st.subheader("❌ Missing Skills")

        if missing:
            for skill in missing:
                st.error(skill.title())
        else:
            st.success("No missing skills.")

    st.write("---")

    st.subheader("📝 Need a Better Resume?")

    if st.button("✨ Go to AI CV Maker", use_container_width=True):
        st.session_state.page = "AI CV Maker"
        st.rerun()


# =====================================================
# RESUME ANALYZER PAGE
# =====================================================
if st.session_state.page == "Resume Analyzer":

    st.header("🔍 Resume Analysis")

    col1, col2 = st.columns(2)

    with col1:
        uploaded_file = st.file_uploader(
            "📤 Upload Resume",
            type=["pdf", "docx", "txt"]
        )

    with col2:
        selected_role = st.selectbox(
            "🎯 Select or Type Job Role",
            job_roles,
            accept_new_options=True,
            placeholder="Select or type custom job role"
        )

        selected_role = selected_role.strip()

    analyze_clicked = st.button("🚀 Analyze Resume", use_container_width=True)

    if uploaded_file and analyze_clicked:

        st.session_state.analysis_result = None

        with st.spinner("Reading Resume..."):
            resume_text = extract_text(uploaded_file)

        if not resume_text or not resume_text.strip():
            st.error("❌ Unable to read resume text.")
            st.stop()

        resume_hash = make_resume_hash(resume_text, selected_role)
        cached_result = get_cached_result(resume_hash)

        if cached_result:

            ai = cached_result.get("ai", {})
            score = cached_result.get("score", 0)
            matched = cached_result.get("matched", [])
            missing = cached_result.get("missing", [])

            st.session_state.analysis_result = {
                "ai": ai,
                "score": score,
                "matched": matched,
                "missing": missing
            }

            st.rerun()

        else:

            current_credits = get_credits(st.session_state.user_email)

            if current_credits < 20:
                 st.error("❌ You need at least 20 credits to analyze a resume.")
                 st.stop()
            
            ok, msg = deduct_credit(st.session_state.user_email, 20)

            if not ok:
                st.error(msg)
                st.stop()

            with st.spinner("Analyzing resume..."):
                slot = wait_for_gemini_slot()

                if not slot:
                    st.error("AI server busy. Try again later.")
                    st.stop()

                ai = validate_resume(resume_text)

            if ai.get("is_resume") is False:
                st.error("Invalid Resume Detected")
                st.warning(ai.get("reason", "This file does not look like a resume."))
                st.stop()

            if ai.get("is_resume") is None:
                ai = {
                    "is_resume": True,
                    "candidate_name": "Unknown",
                    "experience_level": "Unknown",
                    "confidence": 0,
                    "summary": "AI validation unavailable."
                }

            skills = extract_skills(resume_text)
            required_skills = roles_data.get(selected_role, [])

            keyword_score, keyword_matched, keyword_missing = match_skills(
                skills,
                required_skills
            )

            with st.spinner("Running AI-powered resume analysis..."):
                slot = wait_for_gemini_slot()

                if not slot:
                    st.error("AI server busy. Try again later.")
                    st.stop()

                ai_analysis = analyze_resume_ai(
                    resume_text,
                    selected_role,
                    required_skills
                )

            if ai_analysis:
                score = ai_analysis.get("ai_score", keyword_score)
                matched = ai_analysis.get("matched_skills", keyword_matched)
                missing = ai_analysis.get("missing_skills", keyword_missing)

                ai["candidate_level"] = ai_analysis.get("candidate_level", "Unknown")
                ai["role_fit"] = ai_analysis.get("role_fit", "Unknown")
                ai["ai_summary"] = ai_analysis.get("short_summary", "")
                ai["strengths"] = ai_analysis.get("strengths", [])
                ai["weaknesses"] = ai_analysis.get("weaknesses", [])
                ai["improvement_suggestions"] = ai_analysis.get("improvement_suggestions", [])
            else:
                score = keyword_score
                matched = keyword_matched
                missing = keyword_missing

                ai["ai_summary"] = "AI analysis unavailable. Showing keyword-based ATS result."
                ai["strengths"] = []
                ai["weaknesses"] = []
                ai["improvement_suggestions"] = []

            save_cached_result(
                resume_hash=resume_hash,
                selected_role=selected_role,
                ai=ai,
                score=score,
                matched=matched,
                missing=missing
            )

            st.session_state.analysis_result = {
                "ai": ai,
                "score": score,
                "matched": matched,
                "missing": missing
            }

            st.rerun()

    elif not uploaded_file:
        st.session_state.analysis_result = None

        st.info("📤 Upload a resume to begin.")

        st.write("---")

        st.subheader("Don’t have a resume?")

        if st.button("✨ Create Resume with AI", use_container_width=True):
            st.session_state.page = "AI CV Maker"
            st.rerun()

    if st.session_state.analysis_result:
        result = st.session_state.analysis_result

        show_analysis_result(
            result["ai"],
            result["score"],
            result["matched"],
            result["missing"]
        )


# =====================================================
# CV MAKER PAGE
# =====================================================
elif st.session_state.page == "AI CV Maker":

    st.header("✨ AI CV Maker")
    st.caption("Generate professional ATS-friendly resumes using AI.")

    user = get_user(st.session_state.user_email)
    default_name = user.get("name", "") if user else ""

    with st.form("cv_form"):

        name = st.text_input("Full Name", value=default_name)

        email = st.text_input(
            "Email",
            value=st.session_state.user_email
        )

        phone = st.text_input("Phone")

        target_role = st.selectbox(
            "Target Role (Optional)",
            [""] + job_roles,
            accept_new_options=True,
            placeholder="Optional — Select or type target role"
        )

        target_role = target_role.strip() if target_role else ""

        target_role = target_role.strip()

        st.markdown("### Extra Resume Sections")

        add_photo_space = st.radio(
            "Add space for Photo?",
            ["No", "Yes"],
            horizontal=True,
            key="add_photo_space"
        )

        add_certificate_space = st.radio(
            "Add space for Certificates?",
            ["No", "Yes"],
            horizontal=True,
            key="add_certificate_space"
        )

        education = st.text_area("Education")

        skills = st.text_area("Skills")

        experience = st.text_area("Experience")

        projects = st.text_area("Projects")

        submitted = st.form_submit_button(
            "🚀 Generate Resume",
            use_container_width=True
        )

    if submitted:

        if not name or not email or not skills or not education:
            st.error("Please fill at least Name, Email, Education and Skills.")
            st.stop()

        current_credits = get_credits(st.session_state.user_email)

        if current_credits < 50:
            st.error("❌ You need at least 50 credits to generate a resume.")
            st.stop()

        try:
            with st.spinner("Generating Resume using AI..."):
                wait_for_gemini_slot()

                generated_resume = generate_resume(
                    name=name,
                    email=email,
                    phone=phone,
                    education=education,
                    skills=skills,
                    experience=experience,
                    projects=projects,
                    target_role=target_role,
                    add_photo_space=add_photo_space,
                    add_certificate_space=add_certificate_space
                )

            if not generated_resume or "Gemini AI is currently unavailable" in generated_resume:
                st.error("Generation failed , Try gain later")
                st.stop()

            ok, msg = deduct_credit(st.session_state.user_email, 50)

            if not ok:
                st.error(msg)
                st.stop()

            st.session_state.generated_resume = generated_resume
            st.session_state.generated_name = name

        except Exception:
            st.error("Generation failed , Try gain later")
            st.stop()

    if "generated_resume" in st.session_state:

        generated_resume = st.session_state.generated_resume
        resume_name = st.session_state.generated_name

        st.success("Resume Generated Successfully")

        st.subheader("📄 Resume Preview")
        show_resume_preview(
            generated_resume,
            add_photo_space=st.session_state.get("add_photo_space", "No"),
            add_certificate_space=st.session_state.get("add_certificate_space", "No")
        )

        with st.expander("View Raw Text Resume"):
            st.text_area(
                "Raw Resume Text",
                generated_resume,
                height=350
            )

        pdf_data = create_pdf(generated_resume)
        docx_data = create_docx(generated_resume)

        d1, d2, d3 = st.columns(3)

        with d1:
            st.download_button(
                "⬇️ Download TXT",
                data=generated_resume,
                file_name=f"{resume_name.replace(' ', '_')}_resume.txt",
                mime="text/plain",
                use_container_width=True
            )

        with d2:
            st.download_button(
                "📄 Download PDF",
                data=pdf_data,
                file_name=f"{resume_name.replace(' ', '_')}_resume.pdf",
                mime="application/pdf",
                use_container_width=True
            )

        with d3:
            st.download_button(
                "📝 Download DOCX",
                data=docx_data,
                file_name=f"{resume_name.replace(' ', '_')}_resume.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

        st.write("---")

        if st.button("🔍 Analyze This Resume", use_container_width=True):
            st.session_state.page = "Resume Analyzer"
            st.rerun()

# =====================================================
# INTERVIEW PREP PAGE
# =====================================================

if st.session_state.page == "Interview Prep":

    st.markdown("## 🎯 Interview Preparation & Skill Test")

    interview_role = st.selectbox(
        "Choose Job Role",
        job_roles,
        key="interview_role"
    )

    progress = get_user_progress(
        st.session_state.user_email,
        interview_role
    )

    unlocked_level = progress.get("unlocked_level", 1)
    stars = progress.get("stars", 0)
    best_score = progress.get("best_score", 0)

    st.info(f"⭐ Stars: {stars} | 🏆 Best Score: {best_score}%")

    st.markdown("""
    <style>
    .level-card {
        border: 1px solid #333;
        border-radius: 18px;
        padding: 18px;
        text-align: center;
        background: #111111;
        min-height: 300px;
        margin-bottom: 12px;
    }

    .level-title {
        color: #00ff4c;
        font-size: 26px;
        font-weight: 900;
        margin-bottom: 18px;
    }

    .level-info {
        background: #1A1F2E;
        color: white;
        padding: 10px;
        border-radius: 10px;
        margin: 10px 0;
        font-size: 15px;
    }
    </style>
    """, unsafe_allow_html=True)

    # =====================================================
    # LEVELS SECTION
    # =====================================================

    if "interview_questions" not in st.session_state:

        st.markdown("## Levels")

        cols = st.columns(5)

        for level in range(1, 6):

            with cols[level - 1]:

                details = LEVEL_DETAILS[level]

                with st.container(border=True):

                    st.markdown(f"### Level {level}")

                    st.info(details["difficulty"])

                    st.info(f'{details["questions"]} Questions')

                    st.info(f'⏱ {details["time"] // 60} mins')

                    st.info("🪙 50 Credits")

                if level <= unlocked_level:

                    if st.button(
                        f"Start Level {level}",
                        key=f"start_interview_level_{level}",
                        use_container_width=True
                    ):

                        success, message, questions = generate_interview_questions(
                            st.session_state.user_email,
                            interview_role,
                            level
                        )

                        if success:
                            st.session_state.interview_questions = questions
                            st.session_state.interview_level = level
                            st.session_state.interview_role_active = interview_role
                            st.session_state.interview_start_time = time.time()
                            st.session_state.interview_time_limit = LEVEL_DETAILS[level]["time"]
                            st.rerun()
                        else:
                            st.error(message)

                else:

                    st.button(
                        "🔒 Locked",
                        key=f"locked_interview_level_{level}",
                        disabled=True,
                        use_container_width=True
                    )
    
    # =====================================================
    # RESULT POPUP
    # =====================================================

    @st.dialog("🎉 Level Completed")
    def result_popup():

        st.markdown("""
        <style>

        /* POPUP BOX */
        div[data-testid="stDialog"] section[role="dialog"]{
            background: #000000 !important;
            border: 2px solid #00ff4c !important;
            border-radius: 22px !important;

            box-shadow:
                0 0 10px #00ff4c,
                0 0 20px #00ff4c,
                0 0 40px rgba(0,255,76,0.7),
                0 0 80px rgba(0,255,76,0.4) !important;
        }

        </style>
        """, unsafe_allow_html=True)

        st.success("🎉 Congrats!")

        st.markdown(
            f"## You completed Level {st.session_state.interview_level}"
        )

        st.info(
            f"✅ Correct Answers: {st.session_state.correct_answers}"
        )

        st.error(
            f"❌ Wrong Answers: {st.session_state.wrong_answers}"
        )

        st.success(
            f"🏆 Final Score: {st.session_state.final_score}%"
        )

        if st.button(
            "Continue",
            use_container_width=True,
            key="continue_result_popup"
        ):
            st.session_state.show_result_popup = False

            if "interview_questions" in st.session_state:
                del st.session_state["interview_questions"]

            st.rerun()


    if st.session_state.get("show_result_popup"):
        result_popup()
        st.stop()

    # =====================================================
    # ACTIVE EXAM
    # =====================================================

    if "interview_questions" in st.session_state:
        # TIMER
        if not st.session_state.get("submitting_interview", False):
            st_autorefresh(interval=5000, key="interview_timer_refresh")

        elapsed = int(time.time() - st.session_state.interview_start_time)
        remaining = st.session_state.interview_time_limit - elapsed

        if remaining <= 0:
            remaining = 0
            st.session_state.auto_submit_interview = True

        mins = remaining // 60
        secs = remaining % 60

        st.warning(f"⏱ Time Left: {mins:02d}:{secs:02d}")


        st.markdown("---")
        st.subheader(f"📝 Level {st.session_state.interview_level} Test")

        user_answers = {}
        questions = st.session_state.interview_questions

        for i, q in enumerate(questions):

            st.markdown(f"### Q{i + 1}. {q['question']}")

            answer = st.radio(
                "Choose Answer",
                q["options"],
                key=f"interview_answer_{i}",
                index=None
            )

            user_answers[str(i)] = answer

        submit_clicked = st.button(
            "🚀 Submit Test",
            use_container_width=True,
            key="submit_interview_test"
        )

        if submit_clicked or st.session_state.get("auto_submit_interview", False):

            st.session_state.submitting_interview = True

            score, results = check_answers(
                questions,
                user_answers
            )

            stars_earned, new_level = update_user_progress(
                st.session_state.user_email,
                st.session_state.interview_role_active,
                st.session_state.interview_level,
                score
            )

            st.session_state.show_result_popup = True
            st.session_state.final_score = score

            st.session_state.correct_answers = sum(
                1 for r in results if r["is_correct"]
            )

            st.session_state.wrong_answers = sum(
                1 for r in results if not r["is_correct"]
            )

            st.session_state.result_details = results

            st.session_state.auto_submit_interview = False
            st.session_state.submitting_interview = False

            st.rerun()
